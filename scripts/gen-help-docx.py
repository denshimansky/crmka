"""Выгружает app/src/lib/page-help-content.ts в простой .docx.

Формат: для каждой страницы — заголовок (название), затем подзаголовок,
ключ страницы, и далее разделы (heading/text/bullets).
"""

import json
import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "app" / "src" / "lib" / "page-help-content.ts"
OUT = ROOT / "page-help-content.docx"


def extract_object_literal(ts_source: str) -> str:
    """Берёт всё после `= ` и до конца файла — это сам object literal."""
    m = re.search(
        r"export const pageHelpContent[^=]*=\s*(\{[\s\S]*\})\s*\Z",
        ts_source,
    )
    if not m:
        raise SystemExit("Не нашёл объект pageHelpContent в TS-файле")
    return m.group(1)


def js_object_to_json(obj_literal: str) -> dict:
    """Просим Node превратить JS object literal в JSON — единственный
    надёжный способ, т.к. в нём незакавыченные ключи, и т.п."""
    tmp = ROOT / "scripts" / "__dump_help.js"
    tmp.write_text(
        f"const obj = {obj_literal};\nprocess.stdout.write(JSON.stringify(obj));\n",
        encoding="utf-8",
    )
    try:
        res = subprocess.run(
            ["node", str(tmp)],
            capture_output=True,
            text=True,
            encoding="utf-8",
            check=True,
        )
        return json.loads(res.stdout)
    finally:
        tmp.unlink(missing_ok=True)


def build_doc(data: dict) -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading("Справка CRM — тексты для подсказок «?»", level=0)
    p = doc.add_paragraph()
    p.add_run(
        "Источник: app/src/lib/page-help-content.ts. "
        f"Всего страниц: {len(data)}."
    ).italic = True

    for key, content in data.items():
        doc.add_heading(content.get("title", key), level=1)
        if content.get("subtitle"):
            sp = doc.add_paragraph()
            sp.add_run(content["subtitle"]).italic = True
        kp = doc.add_paragraph()
        kp.add_run(f"Ключ страницы: {key}").italic = True

        for section in content.get("sections", []):
            if section.get("heading"):
                doc.add_heading(section["heading"], level=2)
            if section.get("text"):
                doc.add_paragraph(section["text"])
            if section.get("items"):
                for item in section["items"]:
                    doc.add_paragraph(item, style="List Bullet")
    return doc


def main() -> int:
    if not SRC.exists():
        print(f"Нет файла: {SRC}", file=sys.stderr)
        return 1
    obj_literal = extract_object_literal(SRC.read_text(encoding="utf-8"))
    data = js_object_to_json(obj_literal)
    doc = build_doc(data)
    doc.save(OUT)
    print(f"Сохранено: {OUT}  ({len(data)} страниц)")
    return 0


if __name__ == "__main__":
    sys.exit(main())

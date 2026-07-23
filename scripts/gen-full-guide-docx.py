# -*- coding: utf-8 -*-
"""Собирает полное руководство пользователя CRMки в .docx.

Вход: JSON с главами: {"chapters": [{"title": str, "guide": str(markdown),
"discrepancies": [...], "corrections": [...]}]}
Выход:
  - crmka-user-guide.docx — само руководство
  - crmka-user-guide.md   — исходный сводный markdown (для правок)
  - help-content-audit-<дата>.md — сводка расхождений подсказок с кодом

Использование: python scripts/gen-full-guide-docx.py <chapters.json>
"""

import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUT_DOCX = ROOT / "crmka-user-guide.docx"
OUT_MD = ROOT / "crmka-user-guide.md"
OUT_AUDIT = ROOT / "help-content-audit-2026-07-19.md"

BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
CODE_RE = re.compile(r"`([^`]*)`")


def add_runs(paragraph, text: str, italic: bool = False):
    """Разбирает **жирный** и `код` внутри строки и добавляет runs."""
    text = CODE_RE.sub(r"\1", text)
    pos = 0
    for m in BOLD_RE.finditer(text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos:m.start()])
            r.italic = italic
        r = paragraph.add_run(m.group(1))
        r.bold = True
        r.italic = italic
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        r.italic = italic


def render_markdown(doc: Document, md: str, first_chapter: bool):
    lines = md.split("\n")
    i = 0
    first_heading_done = False
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()
        i += 1
        if not stripped:
            continue

        # Заголовки
        m = re.match(r"^(#{1,5})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            if level == 1:
                if not first_chapter or first_heading_done:
                    doc.add_page_break()
                first_heading_done = True
                doc.add_heading(text, level=1)
            else:
                doc.add_heading(text, level=min(level, 4))
            continue

        # Списки (маркированные и нумерованные, с вложенностью по отступу)
        indent = len(raw) - len(raw.lstrip(" "))
        bullet = re.match(r"^[-*•]\s+(.*)$", stripped)
        numbered = re.match(r"^\d+[.)]\s+(.*)$", stripped)
        if bullet or numbered:
            content = (bullet or numbered).group(1)
            level2 = indent >= 2
            if bullet:
                style = "List Bullet 2" if level2 else "List Bullet"
            else:
                style = "List Number 2" if level2 else "List Number"
            try:
                p = doc.add_paragraph(style=style)
            except KeyError:
                p = doc.add_paragraph(style="List Bullet")
            add_runs(p, content)
            continue

        # Строка «Где найти: …» — курсивом
        if stripped.startswith("Где найти:"):
            p = doc.add_paragraph()
            add_runs(p, stripped, italic=True)
            continue

        # Цитата
        if stripped.startswith(">"):
            p = doc.add_paragraph()
            add_runs(p, stripped.lstrip("> "), italic=True)
            continue

        # Обычный абзац
        p = doc.add_paragraph()
        add_runs(p, stripped)


def build(chapters):
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # Титульный лист
    for _ in range(6):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("CRMка")
    r.font.size = Pt(40)
    r.bold = True
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("Полное руководство пользователя")
    r.font.size = Pt(20)
    d = doc.add_paragraph()
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = d.add_run("Версия от 19 июля 2026 года")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_page_break()

    # Оглавление (статичный список; для навигации в Word есть панель «Заголовки»)
    doc.add_heading("Содержание", level=1)
    for ch in chapters:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(ch["title"])
        r.bold = True
        for line in ch["guide"].split("\n"):
            m = re.match(r"^##\s+(.*)$", line.strip())
            if m:
                sub = doc.add_paragraph(style="List Bullet 2")
                sub.add_run(m.group(1).strip())

    first = True
    for ch in chapters:
        render_markdown(doc, ch["guide"], first_chapter=False)
        first = False

    return doc


def build_audit(chapters):
    out = [
        "# Аудит актуальности подсказок «?» — 19.07.2026",
        "",
        "Сверка текстов app/src/lib/page-help-content.ts с реальным кодом страниц.",
        "",
    ]
    total = 0
    for ch in chapters:
        rows = ch.get("discrepancies") or []
        if not rows:
            continue
        out.append(f"## {ch['title']}")
        out.append("")
        for d in rows:
            sev = d.get("severity") or "—"
            out.append(f"- **[{sev}]** `{d.get('pageKey', '?')}` — {d.get('issue', '')}")
            total += 1
        out.append("")
    out.insert(3, f"Всего расхождений: {total}")
    out.insert(4, "")
    return "\n".join(out), total


def main():
    src = Path(sys.argv[1])
    data = json.loads(src.read_text(encoding="utf-8"))
    chapters = data["chapters"]

    md = "\n\n".join(ch["guide"].strip() for ch in chapters)
    OUT_MD.write_text(md, encoding="utf-8")

    audit, total = build_audit(chapters)
    OUT_AUDIT.write_text(audit, encoding="utf-8")

    doc = build(chapters)
    doc.save(OUT_DOCX)
    words = len(md.split())
    print(f"OK: {OUT_DOCX.name} ({len(chapters)} глав, ~{words} слов); аудит: {total} расхождений")


if __name__ == "__main__":
    main()

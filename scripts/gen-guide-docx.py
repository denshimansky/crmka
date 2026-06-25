"""Конвертирует docs/getting-started-guide.md в .docx (Word).

Поддерживает подмножество Markdown, которое используется в руководстве:
заголовки #..####, **жирный**, *курсив*, `код`, маркированные/нумерованные
списки (с одним уровнем вложенности), таблицы, чек-листы «- [ ]» и
горизонтальные разделители «---».

Запуск:  python scripts/gen-guide-docx.py
Требует: pip install python-docx
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "docs" / "getting-started-guide.md"
OUT = ROOT / "getting-started-guide.docx"

SENT = chr(1)  # сентинел для экранированной звёздочки «\*»

INLINE_RE = re.compile(r"\*\*([^*]+)\*\*|`([^`]+)`|\*([^*]+)\*")
H_RE = re.compile(r"^(#{1,4})\s+(.*)$")
HR_RE = re.compile(r"^---+\s*$")
CHECK_RE = re.compile(r"^\s*-\s+\[([ xX])\]\s+(.*)$")
BULLET_RE = re.compile(r"^\s*-\s+")
ORDERED_RE = re.compile(r"^\s*\d+\.\s+")
SUBBULLET_RE = re.compile(r"^\s{2,}-\s+")
SEP_CELL_RE = re.compile(r"^:?-+:?$")


def strip_md(text: str) -> str:
    return text.replace("**", "").replace("`", "").replace("\\*", "*")


def add_runs(paragraph, text: str, bold_all: bool = False) -> None:
    text = text.replace("\\*", SENT)
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos:m.start()].replace(SENT, "*"))
            r.bold = bold_all
        if m.group(1) is not None:
            r = paragraph.add_run(m.group(1).replace(SENT, "*"))
            r.bold = True
        elif m.group(2) is not None:
            r = paragraph.add_run(m.group(2).replace(SENT, "*"))
            r.font.name = "Consolas"
        elif m.group(3) is not None:
            r = paragraph.add_run(m.group(3).replace(SENT, "*"))
            r.italic = True
            r.bold = bold_all
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:].replace(SENT, "*"))
        r.bold = bold_all


def parse_table(lines: list[str]):
    parsed = [
        [c.strip() for c in ln.strip().strip("|").split("|")] for ln in lines
    ]
    sep = None
    for idx, row in enumerate(parsed):
        if row and all(SEP_CELL_RE.match(c) for c in row):
            sep = idx
            break
    header = parsed[sep - 1] if sep and sep > 0 else None
    body = [
        row
        for idx, row in enumerate(parsed)
        if idx != sep and not (sep and idx == sep - 1)
    ]
    return header, body


def add_table(doc: Document, lines: list[str]) -> None:
    header, body = parse_table(lines)
    ncols = max((len(r) for r in [*(([header] if header else [])), *body]), default=1)
    table = doc.add_table(rows=0, cols=ncols)
    table.style = "Table Grid"
    if header:
        cells = table.add_row().cells
        for ci in range(ncols):
            txt = header[ci] if ci < len(header) else ""
            add_runs(cells[ci].paragraphs[0], txt, bold_all=True)
    for row in body:
        cells = table.add_row().cells
        for ci in range(ncols):
            txt = row[ci] if ci < len(row) else ""
            add_runs(cells[ci].paragraphs[0], txt)
    doc.add_paragraph()


def build(md: str) -> Document:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    lines = md.replace("\r\n", "\n").split("\n")
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        if not line.strip():
            i += 1
            continue

        if HR_RE.match(line.strip()):
            doc.add_paragraph()
            i += 1
            continue

        h = H_RE.match(line)
        if h:
            level = len(h.group(1))
            txt = strip_md(h.group(2).strip())
            doc.add_heading(txt, level=0 if level == 1 else level - 1)
            i += 1
            continue

        if line.lstrip().startswith("|"):
            tbl: list[str] = []
            while i < n and lines[i].lstrip().startswith("|"):
                tbl.append(lines[i])
                i += 1
            add_table(doc, tbl)
            continue

        if CHECK_RE.match(line):
            while i < n and CHECK_RE.match(lines[i]):
                m = CHECK_RE.match(lines[i])
                p = doc.add_paragraph(style="List Bullet")
                p.add_run("☑ " if m.group(1).lower() == "x" else "☐ ")
                add_runs(p, m.group(2))
                i += 1
            continue

        if BULLET_RE.match(line):
            while i < n and BULLET_RE.match(lines[i]) and not CHECK_RE.match(lines[i]):
                ln = lines[i]
                indent = len(ln) - len(ln.lstrip())
                style = "List Bullet 2" if indent >= 2 else "List Bullet"
                p = doc.add_paragraph(style=style)
                add_runs(p, BULLET_RE.sub("", ln))
                i += 1
            continue

        if ORDERED_RE.match(line):
            while i < n and (ORDERED_RE.match(lines[i]) or SUBBULLET_RE.match(lines[i])):
                ln = lines[i]
                if ORDERED_RE.match(ln):
                    p = doc.add_paragraph(style="List Number")
                    add_runs(p, ORDERED_RE.sub("", ln))
                else:
                    p = doc.add_paragraph(style="List Bullet 2")
                    add_runs(p, BULLET_RE.sub("", ln))
                i += 1
            continue

        # одиночная строка — абзац (в т.ч. callout «⚠️»)
        p = doc.add_paragraph()
        add_runs(p, line)
        i += 1

    return doc


def main() -> int:
    if not SRC.exists():
        print(f"Нет файла: {SRC}", file=sys.stderr)
        return 1
    md = SRC.read_text(encoding="utf-8-sig")  # utf-8-sig снимает BOM, если есть
    doc = build(md)
    doc.save(OUT)
    print(f"Сохранено: {OUT}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
